# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_md_to_docx(md_file, docx_file):
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(10.5)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)

        # 分隔线
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)

        # 表格
        elif line.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            table_lines = [line]
            i += 1
            # 跳过分隔行
            if lines[i].strip().startswith('|') and '---' in lines[i]:
                i += 1
            # 收集表格行
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            i -= 1

            # 解析表格
            rows = []
            for tline in table_lines:
                cells = [c.strip() for c in tline.split('|')[1:-1]]
                rows.append(cells)

            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = cell_data

        # 列表
        elif line.startswith('- ') or re.match(r'^\d+\. ', line):
            text = re.sub(r'^[-\d]+\.\s*', '', line)
            doc.add_paragraph(text, style='List Bullet')

        # 加粗文本段落
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True

        # 普通段落
        elif line.strip():
            # 处理行内加粗
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.strip('*'))
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    doc.save(docx_file)
    print(f"已生成 Word 文档: {docx_file}")

if __name__ == '__main__':
    md_file = 'AI Gamebot平衡性测试方案_MMO_MOBA通用版.md'
    docx_file = 'AI Gamebot平衡性测试方案_MMO_MOBA通用版.docx'
    parse_md_to_docx(md_file, docx_file)
