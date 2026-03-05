# -*- coding: utf-8 -*-
"""
MD -> Word 转换工具（支持多模板）

用法:
    python md2word.py 输入.md 输出.docx --type 测试报告
    python md2word.py 输入.md 输出.docx --type 测试大纲
    python md2word.py 输入.md 输出.docx --type 工作总结
    python md2word.py 输入.md 输出.docx --template 自定义模板.docx

文档类型与模板对应:
    测试报告  → 测试报告_模板.docx  （左2.5/右2.0，Normal正文）
    测试大纲  → 测试大纲_模板.docx  （左2.5/右2.0，Normal正文）
    工作总结  → 工作总结报告_模板.docx（左右3.2，Body Text First Indent 2正文）
    技术总结  → 技术总结报告_模板.doc（.doc格式，需先转换为.docx后用--template指定）

Markdown 格式约定:
    # 一级标题   → Heading 1
    ## 二级标题  → Heading 2
    ### 三级标题 → Heading 3
    #### 四级标题 → Heading 4
    普通段落     → 正文样式（因模板而异）
    - 列表项     → List Paragraph
    | 表格 |     → Table Grid（表头行自动加粗）
    **粗体**     → 内联加粗
    *斜体*       → 内联斜体
    `代码`       → 内联等宽
    ---          → 分页符
"""

import sys
import re
import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 模板库路径 ─────────────────────────────────────────────────────────────────
_TPLLIB = Path(r'E:\CC\模板库\测试大纲、测试报告、技术总结、工作总结模板')

# type → (模板文件名, 正文样式名)
TEMPLATES = {
    '测试报告': (_TPLLIB / '测试报告_模板.docx',         'Normal'),
    '测试大纲': (_TPLLIB / '测试大纲_模板.docx',         'Normal'),
    '工作总结': (_TPLLIB / '工作总结报告_模板.docx',     'Body Text First Indent 2'),
    # 技术总结为 .doc 格式，需手动转为 .docx 后用 --template 指定
}

DEFAULT_TYPE = '测试报告'


# ── 文档清空 ───────────────────────────────────────────────────────────────────
def clear_body(doc):
    """删除 body 中所有内容段落，保留 sectPr（页面设置）。"""
    body = doc.element.body
    to_remove = [c for c in body if c.tag.split('}')[-1] != 'sectPr']
    for child in to_remove:
        body.remove(child)


# ── 内联格式化 ─────────────────────────────────────────────────────────────────
def add_inline(para, text: str):
    """将含 **bold** / *italic* / `code` 的文本按格式添加到段落。"""
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`)')
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = para.add_run(part[3:-3]); run.bold = True; run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2]); run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1]); run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1]); run.font.name = 'Courier New'
        else:
            para.add_run(part)


# ── 表格解析 ───────────────────────────────────────────────────────────────────
def parse_md_table(lines):
    rows = []
    for line in lines:
        s = line.strip()
        if re.match(r'^\|?[\s:|-]+\|?$', s):
            continue
        rows.append([c.strip() for c in s.strip('|').split('|')])
    return (rows[0], rows[1:]) if rows else ([], [])


def build_table(doc, header, data_rows, body_style):
    all_rows = ([header] + data_rows) if header else data_rows
    if not all_rows:
        return
    col_count = max(len(r) for r in all_rows)
    tbl = doc.add_table(rows=len(all_rows), cols=col_count)
    try:
        tbl.style = 'Table Grid'
    except KeyError:
        pass
    for ri, row_data in enumerate(all_rows):
        for ci in range(col_count):
            cell = tbl.rows[ri].cells[ci]
            para = cell.paragraphs[0]
            para.clear()
            add_inline(para, row_data[ci] if ci < len(row_data) else '')
            if ri == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph(style=body_style)


# ── 代码块 ────────────────────────────────────────────────────────────────────
def build_code_block(doc, lines):
    from docx.shared import Pt
    para = doc.add_paragraph(style='Normal Indent')
    run = para.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(10)


# ── 主转换逻辑 ────────────────────────────────────────────────────────────────
def convert(md_path: str, docx_path: str, template_path: str, body_style: str):
    doc = Document(template_path)
    clear_body(doc)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            build_code_block(doc, code_lines)
            continue

        # 分页符
        if line.strip() == '---':
            para = doc.add_paragraph(style=body_style)
            run = para.add_run()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._element.append(br)
            i += 1
            continue

        # 标题 # ~ ####
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            para = doc.add_paragraph(style=f'Heading {len(m.group(1))}')
            add_inline(para, m.group(2).strip())
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 表格
        if line.strip().startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            header, data = parse_md_table(tbl_lines)
            build_table(doc, header, data, body_style)
            continue

        # 无序 / 有序列表
        if re.match(r'^(\s*)[-*+]\s+', line) or re.match(r'^(\s*)\d+[.)]\s+', line):
            text = re.sub(r'^(\s*)[-*+\d][.)]*\s+', '', line)
            para = doc.add_paragraph(style='List Paragraph')
            add_inline(para, text)
            i += 1
            continue

        # 普通正文
        para = doc.add_paragraph(style=body_style)
        add_inline(para, line.strip())
        i += 1

    doc.save(docx_path)
    print(f'[OK] 已生成: {docx_path}')


# ── 入口 ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description='Markdown -> Word 转换（支持多模板）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('input',  help='输入 Markdown 文件路径')
    parser.add_argument('output', help='输出 Word 文件路径（.docx）')

    group = parser.add_mutually_exclusive_group()
    group.add_argument(
        '--type',
        choices=list(TEMPLATES.keys()),
        default=DEFAULT_TYPE,
        help=f'文档类型（默认：{DEFAULT_TYPE}）'
    )
    group.add_argument(
        '--template',
        help='直接指定模板 .docx 文件路径（会使用 Normal 作为正文样式）'
    )
    parser.add_argument(
        '--body-style',
        default=None,
        help='配合 --template 使用，指定正文段落样式名（默认 Normal）'
    )

    args = parser.parse_args()

    # 确定模板路径和正文样式
    if args.template:
        tpl_path = Path(args.template)
        body_style = args.body_style or 'Normal'
    else:
        doc_type = args.type or DEFAULT_TYPE
        tpl_path, body_style = TEMPLATES[doc_type]
        if args.body_style:
            body_style = args.body_style

    if not Path(args.input).exists():
        print(f'[ERR] 找不到输入文件: {args.input}')
        sys.exit(1)
    if not tpl_path.exists():
        print(f'[ERR] 找不到模板文件: {tpl_path}')
        if not args.template:
            print('  技术总结模板为 .doc 格式，请先用 Word 另存为 .docx，再用 --template 参数指定')
        sys.exit(1)

    convert(args.input, args.output, str(tpl_path), body_style)


if __name__ == '__main__':
    main()
